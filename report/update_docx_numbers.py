"""Update the numbers in report_converted.docx from the committed results.

The document's language, pointers and structure are left untouched: table
cells receive new values in their existing format, and prose sentences have
individual numeric strings replaced in place. Replacements preserve run
formatting by writing into the first run of the matched span.

Run:  .venv/bin/python report/update_docx_numbers.py
"""
import json
import sys
from pathlib import Path

import docx
import numpy as np
import pandas as pd

ROOT = Path(__file__).resolve().parent.parent
DATA = ROOT / "data"
DOC = ROOT / "report" / "report_converted.docx"
MINUS = "−"


# ---------------------------------------------------------------- helpers
def set_cell(cell, text):
    """Replace a cell's text, keeping the first run's formatting."""
    if cell.text.strip() == str(text).strip():
        return False
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = str(text)
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.text = str(text)
    return True


def replace_in_paragraph(p, old, new):
    """Replace old with new inside a paragraph, across runs if needed."""
    full = p.text
    if old not in full:
        return 0
    count = 0
    while old in p.text:
        runs = p.runs
        # locate the run span containing the occurrence
        pos = p.text.index(old)
        acc = 0
        start_run = start_off = None
        for i, r in enumerate(runs):
            if acc + len(r.text) > pos:
                start_run, start_off = i, pos - acc
                break
            acc += len(r.text)
        if start_run is None:
            break
        remaining = len(old)
        i = start_run
        # consume the old string across runs
        first = True
        while remaining > 0 and i < len(runs):
            r = runs[i]
            off = start_off if first else 0
            take = min(remaining, len(r.text) - off)
            r.text = r.text[:off] + (new if first else "") + r.text[off + take:]
            remaining -= take
            first = False
            i += 1
        count += 1
    return count


def replace_everywhere(doc, old, new, where="prose"):
    n = 0
    for p in doc.paragraphs:
        n += replace_in_paragraph(p, old, new)
    return n


def fmt(v, spec):
    s = format(v, spec)
    return s.replace("-", MINUS)


# ---------------------------------------------------------------- data
def main():
    doc = docx.Document(DOC)
    a = pd.read_csv(DATA / "evaluation_accuracy.csv")
    tests = pd.read_csv(DATA / "evaluation_tests.csv")
    lam = pd.read_csv(DATA / "evaluation_lambda.csv")
    thr = pd.read_csv(DATA / "evaluation_threshold.csv")
    ov = pd.read_csv(DATA / "evaluation_overlap.csv")
    w = json.loads((DATA / "learned_weights.json").read_text())
    changed = []

    # ---- table 6: demo plans, Rank (V3, learned) column
    t = doc.tables[6]
    acc3 = a[(a.variant == "V3") & (a["mode"] == "learned")].set_index("plan")
    for r in t.rows[1:]:
        demo = r.cells[0].text.strip().replace("\\_", "_")
        if demo in acc3.index:
            if set_cell(r.cells[7], f"{acc3.loc[demo,'rank_gt']}"):
                changed.append(f"T6 {demo} rank")

    # ---- table 10: learned weight column
    t = doc.tables[10]
    names = {"Density": "t_density", "Span demand": "t_span_demand",
             "Irregularity": "t_irregularity", "Off-wall": "t_off_wall",
             "Repetition": "t_repetition"}
    for r in t.rows[1:]:
        key = names.get(r.cells[0].text.strip())
        if key:
            if set_cell(r.cells[2], fmt(w["weights"][key], "+.2f")):
                changed.append(f"T10 {key}")

    # ---- table 11: the accuracy table
    t = doc.tables[11]
    g = a.groupby(["mode", "variant"])
    order = [("algo", "V1"), ("algo", "V2"), ("algo", "V3"), ("algo", "V4"),
             ("learned", "V1"), ("learned", "V2"), ("learned", "V3"),
             ("learned", "V4")]
    data_rows = [r for r in t.rows[1:] if r.cells[1].text.strip()]
    for row, key in zip(data_rows, order):
        s = g.get_group(key)
        vals = [f"{s.rank_gt.mean():.1f}", f"{s.rank_gt.median():.0f}",
                f"{s.top1.mean():.3f}", f"{s.top5.mean():.3f}",
                f"{s.chamfer_m.mean():.3f}", f"{s.iou.mean():.3f}"]
        for ci, v in enumerate(vals, start=1):
            if set_cell(row.cells[ci], v):
                changed.append(f"T11 {key} col{ci}")

    # ---- table 12: paired tests (learned mode, differing pairs)
    t = doc.tables[12]
    mn = {"rank_gt": "Rank", "chamfer_m": "Distance", "iou": "Overlap"}
    sig = tests[(tests["mode"] == "learned")]
    lookup = {(r.pair, mn[r.metric]): r for r in sig.itertuples()}
    kept = []
    for row in t.rows[1:]:
        pair, metric = row.cells[0].text.strip(), row.cells[1].text.strip()
        r = lookup.get((pair, metric))
        if r is None:
            kept.append((pair, metric))
            continue
        vals = [f"{r.mean_a:.3f}", f"{r.mean_b:.3f}",
                f"{r.wilcoxon_p:.3f}", f"{r.sign_p:.3f}"]
        for ci, v in enumerate(vals, start=2):
            if set_cell(row.cells[ci], v):
                changed.append(f"T12 {pair}/{metric} col{ci}")
    if kept:
        print("  T12 rows with no matching result (left as-is):", kept)

    # ---- table 14: lambda
    t = doc.tables[14]
    lg = lam.groupby(["mode", "lambda"])[["rank_gt", "top1", "top5"]].mean()
    mode = None
    for row in t.rows[1:]:
        c0 = row.cells[0].text.strip()
        if "Fixed" in c0:
            mode = "algo"; continue
        if "Learned" in c0:
            mode = "learned"; continue
        if not c0.startswith("λ"):
            continue
        lv = float(c0.split("=")[1])
        r = lg.loc[(mode, lv)]
        vals = [f"{r.rank_gt:.1f}", f"{r.top1:.3f}", f"{r.top5:.3f}"]
        for ci, v in enumerate(vals, start=1):
            if set_cell(row.cells[ci], v):
                changed.append(f"T14 {mode} λ={lv} col{ci}")

    # ---- table 15: threshold
    t = doc.tables[15]
    ok = thr[thr.derivable == True]                                # noqa: E712
    tg = ok.groupby(["mode", "threshold_mm", "variant"])[
        ["rank_gt", "top1", "top5", "chamfer_m", "gt_n_columns"]].mean()
    mode = None
    for row in t.rows[1:]:
        c0 = row.cells[0].text.strip()
        if "Fixed" in c0:
            mode = "algo"; continue
        if "Learned" in c0:
            mode = "learned"; continue
        if not c0.isdigit():
            continue
        th, v = int(c0), row.cells[1].text.strip()
        r = tg.loc[(mode, th, v)]
        vals = [f"{r.gt_n_columns:.0f}", f"{r.rank_gt:.0f}",
                f"{r.top1:.3f}", f"{r.top5:.3f}", f"{r.chamfer_m:.3f}"]
        for ci, val in enumerate(vals, start=2):
            if set_cell(row.cells[ci], val):
                changed.append(f"T15 {mode} {th}/{v} col{ci}")

    # ---- prose replacements: old string -> new string, computed
    L = a[a["mode"] == "learned"].groupby("variant")
    v1, v2, v3, v4 = (L.get_group(k) for k in ("V1", "V2", "V3", "V4"))
    lamL = lam[lam["mode"] == "learned"].groupby("lambda").rank_gt.mean()
    thrL = ok[ok["mode"] == "learned"].groupby("threshold_mm")
    thrV3 = ok[(ok["mode"] == "learned") & (ok.variant == "V3")].groupby("threshold_mm")

    def words(n):
        return {0: "no", 1: "one", 2: "two", 3: "three", 4: "four",
                5: "five", 6: "six", 7: "seven", 8: "eight", 9: "nine",
                10: "ten", 11: "eleven", 12: "twelve"}[int(n)]

    med3 = words(v3.rank_gt.median())
    n_v4_first = int(v4.top1.sum())
    lamJ = lam[lam["mode"] == "learned"].groupby("lambda")[
        ["top1", "top5"]].mean()
    d_t1 = round((lamJ.loc[0.0, "top1"] - lamJ.loc[1.0, "top1"]) * 7)
    d_t5 = round((lamJ.loc[1.0, "top5"] - lamJ.loc[0.0, "top5"]) * 7)

    prose = [
        # abstract and conclusions: the headline median
        ("from about 2,400 to eight", f"from about 2,400 to {med3}"),
        ("from about 2,395 to eight", f"from about 2,395 to {med3}"),
        # 5.5: how many plans the quantum judge recovers first
        ("recovers the arrangement first on no plan",
         f"recovers the arrangement first on {words(n_v4_first)} plan"
         + ("s" if n_v4_first > 1 else "")),
        # lambda narrative: first-place / top-five counts
        ("costs one first-place recovery, and gains one top-five placing",
         f"costs {words(abs(d_t1))} first-place "
         f"recover{'y' if abs(d_t1) == 1 else 'ies'}, and gains "
         f"{words(abs(d_t5))} top-five placing"
         + ("s" if abs(d_t5) > 1 else "")),
        # 5.5 narrative
        ("median rank of 8, recovers it exactly on two of the seven plans, "
         "and averages 0.246 m",
         f"median rank of {v3.rank_gt.median():.0f}, recovers it exactly on "
         f"{words((v3.rank_gt == 1).sum())} of the seven plans, and averages "
         f"{v3.chamfer_m.mean():.3f} m"),
        ("with 0.749 overlap", f"with {v3.iou.mean():.3f} overlap"),
        ("reducing distance from 0.246 to 0.191 m and raising overlap from "
         "0.749 to 0.790",
         f"reducing distance from {v3.chamfer_m.mean():.3f} to "
         f"{v2.chamfer_m.mean():.3f} m and raising overlap from "
         f"{v3.iou.mean():.3f} to {v2.iou.mean():.3f}"),
        # significance
        ("The smallest p-value obtained is 0.125",
         f"The smallest p-value obtained is "
         f"{tests[tests['mode']=='learned'].wilcoxon_p.min():.3f}"),
        ("the smallest p-value is 0.125",
         f"the smallest p-value is "
         f"{tests[tests['mode']=='learned'].wilcoxon_p.min():.3f}"),
        # lambda narrative
        ("from 23.1 to 22.9",
         f"from {lamL.loc[0.0]:.1f} to {lamL.loc[1.0]:.1f}"),
        # threshold narrative
        ("mean rank of about 25 and 23 respectively and distance of 0.235 "
         "and 0.246 m",
         f"mean rank of about {thrV3.get_group(140).rank_gt.mean():.0f} and "
         f"{thrV3.get_group(160).rank_gt.mean():.0f} respectively and "
         f"distance of {thrV3.get_group(140).chamfer_m.mean():.3f} and "
         f"{thrV3.get_group(160).chamfer_m.mean():.3f} m"),
        ("reaching rank 99 at 180 mm and 214 at 200 mm with distance rising "
         "to 0.616 m",
         f"reaching rank {thrV3.get_group(180).rank_gt.mean():.0f} at 180 mm "
         f"and {thrV3.get_group(200).rank_gt.mean():.0f} at 200 mm with "
         f"distance rising to {thrV3.get_group(200).chamfer_m.mean():.3f} m"),
        # overlap correlation
        (f"mean of 0.148 and a range of {MINUS}0.019 to 0.348",
         f"mean of {ov.spearman_penalty_vs_judge.mean():.3f} and a range of "
         f"{fmt(ov.spearman_penalty_vs_judge.min(), '.3f')} to "
         f"{ov.spearman_penalty_vs_judge.max():.3f}"),
        ("mean of 0.148 and a range of -0.019 to 0.348",
         f"mean of {ov.spearman_penalty_vs_judge.mean():.3f} and a range of "
         f"{fmt(ov.spearman_penalty_vs_judge.min(), '.3f')} to "
         f"{ov.spearman_penalty_vs_judge.max():.3f}"),
        ("rank correlation between them, measured across the full candidate "
         "set of each plan, has a mean of 0.148",
         "rank correlation between them, measured across the full candidate "
         f"set of each plan, has a mean of "
         f"{ov.spearman_penalty_vs_judge.mean():.3f}"),
    ]
    for old, new in prose:
        if old == new:
            continue
        n = replace_everywhere(doc, old, new)
        if n:
            changed.append(f"prose: '{old[:48]}...' x{n}")

    doc.save(DOC)
    print(f"\n{len(changed)} changes written to {DOC.name}")
    for c in changed:
        print("  ", c)


if __name__ == "__main__":
    main()
